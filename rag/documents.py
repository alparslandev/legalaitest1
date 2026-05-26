"""Belgeleri okuma ve parcalara (chunk) ayirma."""
import re
from pathlib import Path
from typing import Iterator

from . import config


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    from docx import Document  # python-docx
    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_pdf(path: Path) -> str:
    # PDFium tabanli; gomulu font eslemesi bozuk olsa bile Turkce karakterleri
    # (ozellikle "i") dogru cikarir. Yalnizca metin-tabanli PDF'ler icin; taranmis
    # /goruntu PDF'lerden metin cikmaz (OCR gerekir).
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(str(path))
    pages = []
    try:
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            pages.append(textpage.get_text_range())
            textpage.close()
            page.close()
    finally:
        pdf.close()
    return "\n\n".join(pages)


READERS = {
    ".txt": _read_txt,
    ".md": _read_txt,
    ".docx": _read_docx,
    ".pdf": _read_pdf,
}


def read_document(path: Path) -> str:
    reader = READERS.get(path.suffix.lower())
    if reader is None:
        raise ValueError(f"Desteklenmeyen dosya turu: {path.suffix}")
    return reader(path)


# PDF/web kopyalamasinda gelen, anlam tasimayan satirlar:
_NOISE_PATTERNS = [
    re.compile(r"https?://\S*", re.I),                       # URL'ler
    re.compile(r"^\d{1,2}\.\d{2}\.\d{4}\s+\d{1,2}:\d{2}"),    # tarih-saat basligi: "14.08.2024 15:08 ..."
    re.compile(r"^\d+\s*/\s*\d+$"),                           # sayfa no: "1/15"
    re.compile(r"^[_\s]+$"),                                  # tek basina "_" / bosluk
    re.compile(r"^(document|_?id=)", re.I),                   # bolunmus URL parcalari
]

# Cumle sonu: . ! ? ardindan bosluk ve buyuk harf/rakam (Turkce harfler dahil)
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?])\s+(?=[\"'(\[]?[A-ZÇĞİÖŞÜ0-9])")


def _split_long(text: str, size: int) -> list[str]:
    return [text[i:i + size] for i in range(0, len(text), size)]


def clean_text(text: str) -> str:
    """Gurultu satirlarini (URL, tarih basligi, sayfa no) atar."""
    kept = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if any(p.search(s) for p in _NOISE_PATTERNS):
            continue
        kept.append(s)
    return "\n".join(kept)


def _sentences(text: str) -> list[str]:
    # PDF satir kirilmalari yapaydir; tek \n'leri bosluga cevirip cumlelere ayir.
    flat = re.sub(r"\s*\n\s*", " ", text)
    flat = re.sub(r"\s{2,}", " ", flat).strip()
    return [s.strip() for s in _SENTENCE_SPLIT.split(flat) if s.strip()]


def chunk_text(text: str) -> list[str]:
    """Once gurultuyu temizler, sonra metni cumle butunlugunu koruyarak
    ~CHUNK_TARGET_CHARS boyutunda parcalara boler; parcalar arasi orpusme birakir."""
    sentences = _sentences(clean_text(text))

    chunks: list[str] = []
    current = ""
    for s in sentences:
        if current and len(current) + len(s) + 1 > config.CHUNK_TARGET_CHARS:
            chunks.append(current)
            overlap = current[-config.CHUNK_OVERLAP_CHARS:]
            current = f"{overlap} {s}".strip()
        else:
            current = f"{current} {s}".strip() if current else s
    if current.strip():
        chunks.append(current)

    # Asiri uzun tek cumle guvenligi (cok nadir): sert bol.
    final: list[str] = []
    for c in chunks:
        if len(c) > config.CHUNK_TARGET_CHARS * 1.6:
            final.extend(_split_long(c, config.CHUNK_TARGET_CHARS))
        else:
            final.append(c)
    return final


def iter_document_chunks(documents_dir: Path) -> Iterator[dict]:
    """documents/ altindaki tum desteklenen dosyalari gezip parca uretir."""
    for path in sorted(documents_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in READERS:
            chunks = chunk_text(read_document(path))
            if not chunks:
                print(f"  UYARI: '{path.name}' icinden metin cikarilamadi "
                      f"(taranmis/goruntu PDF olabilir; OCR gerekir).")
            for i, chunk in enumerate(chunks):
                yield {"source": path.name, "chunk_index": i, "text": chunk}
